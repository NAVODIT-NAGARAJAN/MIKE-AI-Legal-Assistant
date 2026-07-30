"""
DOCX Parser tool for the Document Intelligence Agent.

Responsibilities:
- Accept a .docx file path or raw bytes as input.
- Extract text from paragraphs, tables, headers, and footers.
- Preserve document reading order as closely as possible.
- Return a ParsedDocument result compatible with pdf_parser.ParsedDocument.

This module does NOT perform:
- OCR
- Legal analysis
- Entity extraction
- Clause detection
- Risk detection
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Union

import docx
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.utils.logger import get_logger

log = get_logger(__name__)

# ---------------------------------------------------------------------------
# Minimum character count required to consider a document non-empty.
# ---------------------------------------------------------------------------
_EMPTY_DOCUMENT_CHAR_THRESHOLD = 10


# =============================================================================
# Result dataclass — mirrors pdf_parser.ParsedDocument for pipeline compat.
# =============================================================================


@dataclass
class ParsedDocument:
    """
    Structured result returned by the DOCX parser.

    Fields:
        file_name       : Original file name, if available.
        page_count      : Always None for DOCX (page boundaries are dynamic).
        sections        : Per-section extracted text. Index 0 = section 1.
        full_text       : All text joined with double newlines.
        character_count : Total character count of full_text.
        ocr_applied     : Always False — DOCX parser does not use OCR.
        scanned_pages   : Always empty — not applicable for DOCX.
    """

    file_name: Optional[str]
    page_count: Optional[int]
    sections: List[str] = field(default_factory=list)
    full_text: str = ""
    character_count: int = 0
    ocr_applied: bool = False
    scanned_pages: List[int] = field(default_factory=list)


# =============================================================================
# Exceptions
# =============================================================================


class DOCXParserError(Exception):
    """
    Raised when the DOCX parser cannot process the provided document.
    """


class EmptyDocumentError(DOCXParserError):
    """
    Raised when the DOCX file yields no extractable text.
    """


# =============================================================================
# Internal helpers
# =============================================================================


def _open_document(source: Union[str, Path, bytes]) -> tuple[Document, Optional[str]]:
    """
    Open a python-docx Document from a file path or raw bytes.

    Returns:
        (Document, file_name) — file_name is None when source is bytes.

    Raises:
        DOCXParserError: If the source cannot be opened as a DOCX document.
    """
    if isinstance(source, bytes):
        if len(source) == 0:
            raise DOCXParserError("DOCX bytes are empty.")
        try:
            document = Document(io.BytesIO(source))
            return document, None
        except PackageNotFoundError as exc:
            raise DOCXParserError(f"Invalid DOCX bytes: {exc}") from exc
        except Exception as exc:
            raise DOCXParserError(f"Cannot open DOCX from bytes: {exc}") from exc

    file_path = Path(source)

    if not file_path.exists():
        raise DOCXParserError(f"DOCX file not found: {file_path}")

    if not file_path.is_file():
        raise DOCXParserError(f"Path is not a file: {file_path}")

    if file_path.stat().st_size == 0:
        raise DOCXParserError(f"DOCX file is empty: {file_path}")

    if file_path.suffix.lower() not in {".docx", ".docm"}:
        raise DOCXParserError(
            f"Unsupported file extension '{file_path.suffix}'. "
            "Only .docx and .docm files are supported."
        )

    try:
        document = Document(str(file_path))
        return document, file_path.name
    except PackageNotFoundError as exc:
        raise DOCXParserError(
            f"Cannot read DOCX file '{file_path.name}': file is corrupt or not a valid DOCX. "
            f"Detail: {exc}"
        ) from exc
    except Exception as exc:
        raise DOCXParserError(
            f"Unexpected error opening '{file_path.name}': {exc}"
        ) from exc


def _extract_paragraph_text(document: Document) -> List[str]:
    """
    Extract text from all body paragraphs, preserving order.
    Empty paragraphs are skipped.
    """
    lines: List[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return lines


def _extract_table_text(document: Document) -> List[str]:
    """
    Extract text from all tables in the document body.

    Each table is represented as a block of tab-separated row text,
    with rows separated by newlines.
    """
    blocks: List[str] = []
    for table in document.tables:
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append("\t".join(cells))
        if rows:
            blocks.append("\n".join(rows))
    return blocks


def _extract_header_footer_text(document: Document) -> List[str]:
    """
    Extract text from headers and footers across all sections.
    Duplicate header/footer content across sections is deduplicated.
    """
    seen: set[str] = set()
    lines: List[str] = []

    for section in document.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if part is None:
                continue
            for para in part.paragraphs:
                text = para.text.strip()
                if text and text not in seen:
                    seen.add(text)
                    lines.append(text)

    return lines


# =============================================================================
# Public API
# =============================================================================


def parse_docx(source: Union[str, Path, bytes]) -> ParsedDocument:
    """
    Extract text and metadata from a Microsoft Word (.docx) document.

    Extraction order:
        1. Headers and footers (document-level context)
        2. Body paragraphs (main content, reading order preserved)
        3. Tables (tabular data)

    Args:
        source: Absolute file path (str or Path) or raw DOCX bytes.

    Returns:
        ParsedDocument with full_text and metadata.

    Raises:
        DOCXParserError    : On unreadable, invalid, or unsupported input.
        EmptyDocumentError : When no text can be extracted from the document.
    """

    log.info("=" * 60)
    log.info("DOCXParser: starting extraction.")

    document, file_name = _open_document(source)

    log.info("DOCXParser: opened '%s'.", file_name or "<bytes>")

    # ------------------------------------------------------------------
    # Extract content in reading order
    # ------------------------------------------------------------------
    header_footer_lines = _extract_header_footer_text(document)
    paragraph_lines = _extract_paragraph_text(document)
    table_blocks = _extract_table_text(document)

    log.info(
        "DOCXParser: extracted %d header/footer line(s), "
        "%d paragraph(s), %d table block(s).",
        len(header_footer_lines),
        len(paragraph_lines),
        len(table_blocks),
    )

    # ------------------------------------------------------------------
    # Assemble sections for the result
    # ------------------------------------------------------------------
    sections: List[str] = []

    if header_footer_lines:
        sections.append("\n".join(header_footer_lines))

    if paragraph_lines:
        sections.append("\n".join(paragraph_lines))

    if table_blocks:
        sections.extend(table_blocks)

    full_text = "\n\n".join(s for s in sections if s.strip())
    character_count = len(full_text)

    if character_count < _EMPTY_DOCUMENT_CHAR_THRESHOLD:
        raise EmptyDocumentError(
            "The DOCX document yielded no extractable text. "
            "The file may be empty or contain only images."
        )

    log.info(
        "DOCXParser: extraction complete. "
        "Characters: %d | Sections: %d",
        character_count,
        len(sections),
    )
    log.info("=" * 60)

    return ParsedDocument(
        file_name=file_name,
        page_count=None,
        sections=sections,
        full_text=full_text,
        character_count=character_count,
        ocr_applied=False,
        scanned_pages=[],
    )
